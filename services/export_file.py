# -*- coding: utf-8 -*-
# md -> pdf
# md -> docx
import os
import sys
import argparse
from typing import List, Tuple
import markdown
# For DOCX
from docx import Document
from docx.shared import Inches
# For PDF
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage, Table as RLTable
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.pagesizes import letter
from reportlab.lib.enums import TA_LEFT
from reportlab.lib import colors
from reportlab.platypus import TableStyle
from PIL import Image as PILImage
import subprocess
import os
import sys
from typing import Optional
import re # For simple image path extraction from MD

# --- 1. Markdown 解析和内容提取 ---
def extract_content_from_md(md_text: str) -> Tuple[List[str], List[str], List[List[List[str]]]]:
    """
    从 Markdown 文本中提取纯文本段落、图片路径和表格数据。
    注意：这是一个简化的解析器，适用于基本的 MD 格式。
    对于复杂的 MD (如嵌套列表、代码块)，可能需要更强大的库如 `markdown` + HTML 解析。
    """
    lines = md_text.strip().split('\n')
    text_paragraphs: List[str] = []
    image_paths: List[str] = []
    table_data_list: List[List[List[str]]] = []

    i = 0
    current_paragraph_lines: List[str] = []

    while i < len(lines):
        line = lines[i]

        # --- 识别图片 (Markdown 格式: ![alt](path)) ---
        # 使用正则表达式匹配图片链接
        img_match = re.match(r'!\[.*?\]\((.*?)\)', line.strip())
        if img_match:
            img_path = img_match.group(1)
            if current_paragraph_lines:
                text_paragraphs.append('\n'.join(current_paragraph_lines))
                current_paragraph_lines = []
            # 简单验证路径是否存在或是否为网络路径 (这里假设本地路径)
            # 你可以根据需要添加更复杂的验证
            image_paths.append(img_path)
            i += 1
            continue

        # --- 识别表格 (Markdown 格式) ---
        # 表格通常由 | 分隔，并且第二行是 --- 
        # 检查当前行和下一行是否构成表格开头
        if '|' in line and i + 1 < len(lines) and re.match(r'^\s*\|(\s*[-:]+\s*\|)+\s*$', lines[i+1]):
            if current_paragraph_lines:
                text_paragraphs.append('\n'.join(current_paragraph_lines))
                current_paragraph_lines = []

            table_data = []
            # 添加表头 (当前行)
            header_cells = [cell.strip() for cell in line.split('|') if cell.strip()]
            table_data.append(header_cells)
            i += 2 # 跳过表头和分隔行

            # 收集数据行
            while i < len(lines) and '|' in lines[i]:
                data_cells = [cell.strip() for cell in lines[i].split('|') if cell.strip()]
                # 确保列数一致，不足则补空
                while len(data_cells) < len(header_cells):
                    data_cells.append('')
                # 如果列数过多，则截断
                data_cells = data_cells[:len(header_cells)]
                table_data.append(data_cells)
                i += 1
            
            table_data_list.append(table_data)
            continue # 继续主循环，因为 i 已经更新

        # --- 处理文本 ---
        # 空行表示段落分隔
        if not line.strip():
            if current_paragraph_lines:
                text_paragraphs.append('\n'.join(current_paragraph_lines))
                current_paragraph_lines = []
        else:
            current_paragraph_lines.append(line)
        i += 1

    # 处理最后一个段落
    if current_paragraph_lines:
        text_paragraphs.append('\n'.join(current_paragraph_lines))
    print(image_paths)
    return text_paragraphs, image_paths, table_data_list

# --- 2. DOCX 导出函数 ---
def export_docx(
    output_path: str,
    text_paragraphs: List[str],
    image_paths: List[str],
    table_data_list: List[List[List[str]]],
    image_width_inches: float = 4.0
):
    """将解析后的内容导出为 DOCX 文件。"""
    try:
        doc = Document()
        # 确保资源列表长度与段落数量匹配
        # 例如：段落1 -> 图1 -> 表1 -> 段落2 -> ...
        # 所以我们有 N 个段落，最多 N-1 个图片和 N-1 个表格
        num_resources = len(text_paragraphs) - 1 if text_paragraphs else 0
        # 如果资源多于预期，截断；如果少于，用 None 填充
        image_paths_padded = (image_paths + [None] * max(0, num_resources - len(image_paths)))[:num_resources] if num_resources > 0 else []
        table_data_list_padded = (table_data_list + [None] * max(0, num_resources - len(table_data_list)))[:num_resources] if num_resources > 0 else []

        for i, paragraph_text in enumerate(text_paragraphs):
            if paragraph_text.strip():
                doc.add_paragraph(paragraph_text)
            
            # 在段落之后插入对应的资源（除了最后一个段落）
            if i < num_resources:
                # 插入图片
                img_path = image_paths_padded[i]
                if img_path:
                    img_path = os.path.abspath(img_path)  # 确保路径是绝对路径
                    print(img_path)
                    # 简单检查文件是否存在 (对于本地路径)
                    if not (img_path.startswith('http://') or img_path.startswith('https://')):
                        if not os.path.exists(img_path):
                            doc.add_paragraph(f"[警告: 找不到图片 '{img_path}']")
                        else:
                            try:
                                doc.add_paragraph() # 空行
                                doc.add_picture(img_path, width=Inches(image_width_inches))
                            except Exception as e:
                                doc.add_paragraph(f"[图片插入失败: {img_path} - {e}]")
                    else: # 网络图片，python-docx 不直接支持，需要先下载
                         doc.add_paragraph(f"[网络图片: {img_path}]") # 占位符

                # 插入表格
                table_data = table_data_list_padded[i]
                if table_data:
                    try:
                        doc.add_paragraph() # 空行
                        if table_data and len(table_data) > 0 and len(table_data[0]) > 0:
                            num_cols = len(table_data[0])
                            table = doc.add_table(rows=0, cols=num_cols)
                            table.style = 'Table Grid'
                            
                            # 添加表头并加粗
                            hdr_cells = table.add_row().cells
                            for j, cell_value in enumerate(table_data[0]):
                                hdr_cells[j].text = str(cell_value)
                                # 对于表头加粗，可以应用样式或直接操作 run
                                # 这里简单处理，实际可能需要更复杂的样式设置
                            
                            # 添加数据行
                            for row_data in table_data[1:]:
                                row_cells = table.add_row().cells
                                padded_row_data = (row_data + [''] * num_cols)[:num_cols]
                                for j, cell_value in enumerate(padded_row_data):
                                    row_cells[j].text = str(cell_value)
                        else:
                            doc.add_paragraph("[表格数据为空或无效]")
                    except Exception as e:
                        doc.add_paragraph(f"[表格插入失败: {e}]")

        doc.save(output_path)
        print(f"✅ DOCX 文件已保存至: {output_path}")
        return True
    except Exception as e:
        print(f"❌ 生成 DOCX 文件时出错: {e}")
        import traceback
        traceback.print_exc()
        return False




def convert_docx_to_pdf_with_libreoffice(
    input_docx_path: str,
    output_pdf_path: str,
    libreoffice_command: str = "soffice", # 尝试 'soffice' 或 'libreoffice'
    overwrite: bool = True
) -> bool:
    """
    使用 LibreOffice 命令行工具将 DOCX 文件转换为 PDF 文件。

    此方法符合知识库中提到的“回答内容支持文本、已编撰好的 PDF/Word 等格式的文件”的要求，
    并利用系统已安装的 LibreOffice 工具。

    Args:
        input_docx_path (str): 输入的 DOCX 文件路径。
        output_pdf_path (str): 输出的 PDF 文件路径。
        libreoffice_command (str): 调用 LibreOffice 的命令，默认为 'soffice'。
                                   在某些系统上可能需要 'libreoffice'。
        overwrite (bool): 如果输出文件已存在，LibreOffice 通常会覆盖。

    Returns:
        bool: 转换成功返回 True，否则返回 False。
    """
    # 1. 检查输入文件是否存在
    if not os.path.exists(input_docx_path):
        print(f"❌ 错误: 输入的 DOCX 文件 '{input_docx_path}' 不存在。")
        return False

    # 2. 获取并创建输出目录（如果需要）
    output_dir = os.path.dirname(output_pdf_path)
    if output_dir and not os.path.exists(output_dir):
        try:
            os.makedirs(output_dir)
            print(f"📁 已创建输出目录: {output_dir}")
        except OSError as e:
            print(f"❌ 错误: 无法创建输出目录 '{output_dir}': {e}")
            return False

    # 3. 构建 LibreOffice 命令
    # --headless: 无界面模式运行
    # --convert-to pdf: 指定转换目标格式为 PDF
    # --outdir: 指定输出目录
    cmd = [
        libreoffice_command,
        "--headless",             # 无头模式
        "--convert-to", "pdf",    # 转换为 PDF
        "--outdir", output_dir if output_dir else ".", # 输出目录
        input_docx_path           # 输入文件
    ]

    print(f"🔄 正在调用 LibreOffice 命令: {' '.join(cmd)}")

    try:
        # 4. 执行命令
        result = subprocess.run(
            cmd,
            check=True,           # 如果返回码非0，抛出 CalledProcessError
            capture_output=True,  # 捕获 stdout 和 stderr
            text=True,            # 将输出解码为字符串
            timeout=120           # 设置超时时间（秒）
        )
        print(f"✅ LibreOffice 转换命令执行成功。")

        # 5. 验证输出文件是否存在
        if os.path.exists(output_pdf_path):
            print(f"✅ PDF 文件已成功生成: {output_pdf_path}")
            return True
        else:
            print(f"⚠️  LibreOffice 命令执行成功，但未在 '{output_pdf_path}' 找到 PDF 文件。")
            # 可以尝试列出输出目录内容来调试
            if output_dir:
                print(f"   输出目录 '{output_dir}' 的内容:")
                try:
                    for f in os.listdir(output_dir):
                        print(f"     - {f}")
                except OSError:
                    pass
            return False

    except subprocess.CalledProcessError as e:
        print(f"❌ LibreOffice 转换失败 (返回码 {e.returncode}): {e}")
        if e.stdout:
            print(f"   Stdout: {e.stdout}")
        if e.stderr:
            print(f"   Stderr: {e.stderr}")
        return False
    except subprocess.TimeoutExpired:
        print(f"❌ LibreOffice 转换超时 (超过 120 秒)。")
        return False
    except FileNotFoundError:
        print(f"❌ 未找到命令 '{libreoffice_command}'。请确保 LibreOffice 已安装并且命令在 PATH 中。")
        print(f"   您可能需要尝试使用 'libreoffice' 作为命令。")
        return False
    except Exception as e:
        print(f"❌ 调用 LibreOffice 转换时发生未知错误: {e}")
        import traceback
        traceback.print_exc()
        return False

# --- 示例用法 ---
def export2pdf(input_docx, output_pdf):
    # --- 配置 ---
    # 请将下面的路径替换为您实际的 DOCX 文件路径
    # input_docx = "example_output.docx"  # 输入 DOCX 文件
    # output_pdf = "converted_output.pdf" # 输出 PDF 文件
    # 如果 'soffice' 命令不起作用，请尝试 'libreoffice'
    libreoffice_cmd = "soffice" # 或 "libreoffice"
    # --- 配置结束 ---

    if not os.path.exists(input_docx):
        print(f"⚠️  示例输入文件 '{input_docx}' 不存在。请先生成一个 DOCX 文件或修改路径。")
        # 可以选择退出或提示用户
        # sys.exit(1)

    print(f"📄 准备将 '{input_docx}' 转换为 '{output_pdf}'...")
    success = convert_docx_to_pdf_with_libreoffice(
        input_docx_path=input_docx,
        output_pdf_path=output_pdf,
        libreoffice_command=libreoffice_cmd
    )

    if success:
        print(f"\n🎉 转换成功完成!")
    else:
        print(f"\n💥 转换失败。请检查错误信息。")
        # sys.exit(1) # 根据需要决定是否退出

def md2docx(input_md, output_docx):
    print(f"📄 准备将 '{input_md}' 转换为 '{output_docx}'...")
    # --- 读取 Markdown 文件 ---
    try:
        with open(input_md, 'r', encoding='utf-8') as f:
            md_content = f.read()
        print(f"📄 已读取 Markdown 文件: {input_md}")
    except Exception as e:
        print(f"❌ 读取 Markdown 文件时出错: {e}")
        sys.exit(1)
    texts, images, tables = extract_content_from_md(md_content)

    # --- 生成文件 ---
    print("\n💾 正在生成文件...")
    success_docx = export_docx(output_docx, texts, images, tables)
    if success_docx:
        print(f"\n🎉 转换成功完成!"
              "\n   - DOCX: {output_docx}")
    else:
        print(f"\n💥 转换失败。请检查错误信息.")

def md2pdf(input_md, output_pdf):
    # 确定输出文件名
    base_name = os.path.splitext(input_md)[0]
    docx_path = f"{base_name}.docx"
   
    md2docx(input_md, docx_path)
    success_pdf = export2pdf(input_docx=docx_path, output_pdf=output_pdf)

    if success_pdf:
        print(f"\n🎉 转换成功完成!"
              "\n   - PDF:  {output_pdf}")
    else:
        print(f"\n💥 转换失败。请检查错误信息.")

def mdcontent2docx(md_content, output_docx):
    print(f"📄 准备转换为 '{output_docx}'...")
    texts, images, tables = extract_content_from_md(md_content)

    # --- 生成文件 ---
    print("\n💾 正在生成文件...")
    success_docx = export_docx(output_docx, texts, images, tables)
    if success_docx:
        print(f"\n🎉 转换成功完成!"
              "\n   - DOCX: {output_docx}")
    else:
        print(f"\n💥 转换失败。请检查错误信息.")

def mdcontent2pdf(mdcontent, output_pdf):
    # 确定输出文件名
    base_name = os.path.splitext(output_pdf)[0]
    docx_path = f"{base_name}.docx"
   
    mdcontent2docx(mdcontent, docx_path)
    success_pdf = export2pdf(input_docx=docx_path, output_pdf=output_pdf)

    if success_pdf:
        print(f"\n🎉 转换成功完成!"
              "\n   - PDF:  {output_pdf}")
    else:
        print(f"\n💥 转换失败。请检查错误信息.")

def mdcontent2md(mdcontent: str, md_path: str) -> bool:
    """
    将给定的 Markdown 字符串内容保存到指定的文件路径。

    Args:
        mdcontent (str): 包含 Markdown 语法的字符串内容。
        md_path (str): 输出 Markdown 文件的路径。

    Returns:
        bool: 如果文件保存成功返回 True，否则返回 False。
    """
    try:
        # --- 确保输出目录存在 ---
        output_dir = os.path.dirname(md_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)
            print(f"📁 已创建输出目录: {output_dir}")

        # --- 写入文件 ---
        with open(md_path, 'w', encoding='utf-8') as f:
            f.write(mdcontent)
        
        print(f"✅ Markdown 内容已成功保存至: {md_path}")
        return True

    except Exception as e:
        print(f"❌ 保存 Markdown 文件时出错: {e}")
        import traceback
        traceback.print_exc()
        return False
    

# --- 4. 主程序 ---
def main():
    mdcontent = '''
            ### 语义图像构建方法

            #### 1. **固定部件构建流程**
            - **初始化张量**
                创建尺寸为 `width/8 × height/8`、通道数为 `item_class_num` 的灰度图像张量，初始值为 `0.01`（值域 `[0, 1]`），以节省内存。
            - **生成高斯椭圆**
                对每个固定部件的边界框，计算其中心点为均值 `μ`，半宽/半长为 `2σ`（X/Y方向独立），生成峰值为 `1.0` 的二维正态分布（公式见原文第2-4步）。
                *示例图：*
                ![图1：高斯椭圆分布](D:/adavance/tsy/rag4chat/output/test8/auto/images/258b7f7f1bcf4f9204aeb3191f283fc260e4d9d699d66ff3b1a854fc4c882aa2.jpg)        
            - **叠加语义分布**
                将各部件的高斯椭圆叠加至对应类别通道的图像中，重叠区域取相对高值（如图4所示）。

            #### 2. **旋转移动部件构建**
            - 文档提到其构建方法与固定部件**不一致**，但未提供具体步骤，需结合其他技术（如动态边界框调整）进一步研究。

            #### 3. **视觉分析参考**
            - 图像示例：
                ![图2：多通道语义分布](D:/adavance/tsy/rag4chat/output/test8/auto/images/258b7f7f1bcf4f9204aeb3191f283fc260e4d9d699d66ff3b1a854fc4c882aa2.jpg)      
                （注：此图为技术分析示例，包含红色标注的边界框和语义分布区域）'''
    mdcontent2pdf(mdcontent, "./example.pdf")
    # mdcontent2md(mdcontent, "./example0.md")
    # md2docx("./example0.md", "./example0.docx")
    # parser = argparse.ArgumentParser(description="将 Markdown 文件转换为 DOCX 和 PDF。")
    # parser.add_argument("input_md",default="./example.md", help="输入的 Markdown 文件路径")
    # parser.add_argument("-d", "--docx", help="输出的 DOCX 文件路径 (默认: input.md -> input.docx)")
    # parser.add_argument("-p", "--pdf", help="输出的 PDF 文件路径 (默认: input.md -> input.pdf)")
    
    # args = parser.parse_args()

    # input_md_path = args.input_md

    # if not os.path.exists(input_md_path):
    #     print(f"❌ 错误: 输入文件 '{input_md_path}' 不存在。")
    #     sys.exit(1)

    # # 确定输出文件名
    # base_name = os.path.splitext(input_md_path)[0]
    # output_docx_path = args.docx if args.docx else f"{base_name}.docx"
    # output_pdf_path = args.pdf if args.pdf else f"{base_name}.pdf"

    # # --- 读取 Markdown 文件 ---
    # try:
    #     with open(input_md_path, 'r', encoding='utf-8') as f:
    #         md_content = f.read()
    #     print(f"📄 已读取 Markdown 文件: {input_md_path}")
    # except Exception as e:
    #     print(f"❌ 读取 Markdown 文件时出错: {e}")
    #     sys.exit(1)

    # # --- 解析内容 ---
    # print("🔍 正在解析 Markdown 内容...")
    # texts, images, tables = extract_content_from_md(md_content)

    # print(f"  - 解析到文本段落数: {len(texts)}")
    # print(f"  - 解析到图片路径数: {len(images)}")
    # print(f"  - 解析到表格数量: {len(tables)}")

    # # --- 生成文件 ---
    # print("\n💾 正在生成文件...")
    # success_docx = export_docx(output_docx_path, texts, images, tables)
    # success_pdf = export2pdf(input_docx=output_docx_path, output_pdf=output_pdf_path)

    # if success_docx and success_pdf:
    #     print(f"\n🎉 所有文件已成功生成!")
    #     print(f"   - DOCX: {output_docx_path}")
    #     print(f"   - PDF:  {output_pdf_path}")
    # else:
    #     print(f"\n⚠️  部分文件生成失败。")
    #     sys.exit(1)

if __name__ == "__main__":
    main()